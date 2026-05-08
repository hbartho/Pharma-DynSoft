#!/usr/bin/env python3
"""
Script pour générer le diagramme de classes de DynSoft Pharma en format Word
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def create_class_diagram_doc():
    doc = Document()
    
    # Configuration de la page
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # A4 paysage
    section.page_height = Cm(21)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    # Titre principal
    title = doc.add_heading('DynSoft Pharma - Diagramme de Classes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre
    subtitle = doc.add_paragraph('Application de Gestion de Pharmacie')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph('Généré le 2 Janvier 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # =====================================
    # SECTION 1: ENTITÉS PRINCIPALES
    # =====================================
    doc.add_heading('1. Entités Principales', level=1)
    
    # Table pour User
    doc.add_heading('User (Utilisateur)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    user_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('email', 'str', 'Email de connexion'),
        ('first_name', 'str', 'Prénom'),
        ('last_name', 'str', 'Nom'),
        ('employee_code', 'str', 'Code employé (ex: ADM-001)'),
        ('role', 'str', 'admin | pharmacien | caissier'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('is_active', 'bool', 'Statut actif/inactif'),
        ('created_at', 'datetime', 'Date de création'),
    ]
    for attr in user_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # Table pour Product
    doc.add_heading('Product (Produit)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    product_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('name', 'str', 'Nom du produit'),
        ('internal_reference', 'str?', 'Référence interne'),
        ('barcode', 'str?', 'Code-barres'),
        ('description', 'str?', 'Description'),
        ('purchase_price', 'float', 'Prix d\'achat'),
        ('price', 'float', 'Prix de vente'),
        ('stock', 'int', 'Quantité en stock'),
        ('min_stock', 'int', 'Stock minimum (alerte)'),
        ('category_id', 'str?', 'FK → Category'),
        ('unit_id', 'str?', 'FK → Unit'),
        ('expiration_date', 'datetime?', 'Date de péremption'),
        ('is_active', 'bool', 'Statut actif/inactif'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date de création'),
        ('updated_at', 'datetime', 'Date de modification'),
    ]
    for attr in product_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # Table pour Category
    doc.add_heading('Category (Catégorie)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    category_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('name', 'str', 'Nom de la catégorie'),
        ('description', 'str?', 'Description'),
        ('color', 'str', 'Couleur (hex)'),
        ('markup_coefficient', 'float', 'Coefficient de marge'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date de création'),
    ]
    for attr in category_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # Table pour Unit
    doc.add_heading('Unit (Unité de mesure)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    unit_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('name', 'str', 'Nom (Boîte, Flacon, etc.)'),
        ('abbreviation', 'str?', 'Abréviation (BTE, FLC)'),
        ('description', 'str?', 'Description'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date de création'),
    ]
    for attr in unit_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_page_break()
    
    # =====================================
    # SECTION 2: PARTENAIRES COMMERCIAUX
    # =====================================
    doc.add_heading('2. Partenaires Commerciaux', level=1)
    
    # Table pour Supplier
    doc.add_heading('Supplier (Fournisseur)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    supplier_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('name', 'str', 'Nom du fournisseur'),
        ('phone', 'str?', 'Téléphone'),
        ('email', 'str?', 'Email'),
        ('address', 'str?', 'Adresse'),
        ('is_active', 'bool', 'Statut actif/inactif'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date de création'),
        ('updated_at', 'datetime?', 'Date de modification'),
        ('updated_by', 'str?', 'Code employé modificateur'),
    ]
    for attr in supplier_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # Table pour Customer
    doc.add_heading('Customer (Client)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    customer_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('name', 'str', 'Nom du client'),
        ('phone', 'str?', 'Téléphone'),
        ('email', 'str?', 'Email'),
        ('address', 'str?', 'Adresse'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date de création'),
    ]
    for attr in customer_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # =====================================
    # SECTION 3: TRANSACTIONS
    # =====================================
    doc.add_heading('3. Transactions', level=1)
    
    # Table pour Sale
    doc.add_heading('Sale (Vente)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    sale_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('sale_number', 'str?', 'Numéro de vente (VNT-XXXXXXXX)'),
        ('customer_id', 'str?', 'FK → Customer'),
        ('items', 'List[Dict]', 'Liste des articles vendus'),
        ('total', 'float', 'Montant total'),
        ('payment_method', 'str', 'Mode de paiement'),
        ('user_id', 'str?', 'ID de l\'utilisateur'),
        ('employee_code', 'str?', 'Code employé vendeur'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date de vente'),
    ]
    for attr in sale_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # Table pour SaleReturn
    doc.add_heading('SaleReturn (Retour)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    return_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('return_number', 'str?', 'Numéro de retour (RET-XXXXXXXX)'),
        ('sale_id', 'str', 'FK → Sale'),
        ('sale_number', 'str?', 'Numéro de la vente originale'),
        ('items', 'List[Dict]', 'Articles retournés'),
        ('total_refund', 'float', 'Montant remboursé'),
        ('reason', 'str?', 'Motif du retour'),
        ('user_id', 'str', 'ID de l\'utilisateur'),
        ('employee_code', 'str?', 'Code employé'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date du retour'),
    ]
    for attr in return_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_page_break()
    
    # =====================================
    # SECTION 4: APPROVISIONNEMENTS
    # =====================================
    doc.add_heading('4. Approvisionnements', level=1)
    
    # Table pour Supply
    doc.add_heading('Supply (Approvisionnement)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    supply_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('supply_date', 'datetime', 'Date d\'approvisionnement'),
        ('is_validated', 'bool', 'Statut de validation'),
        ('validated_at', 'datetime?', 'Date de validation'),
        ('validated_by', 'str?', 'Code employé validateur'),
        ('supplier_id', 'str?', 'FK → Supplier'),
        ('supplier_name', 'str?', 'Nom fournisseur (dénormalisé)'),
        ('total_amount', 'float', 'Montant total'),
        ('purchase_order_ref', 'str?', 'Réf. bon de commande'),
        ('delivery_note_number', 'str?', 'N° bon de livraison'),
        ('invoice_number', 'str?', 'N° facture'),
        ('is_credit_note', 'bool', 'Est un avoir'),
        ('notes', 'str?', 'Notes'),
        ('items', 'List[SupplyItem]', 'Détails des produits'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date de création'),
        ('created_by', 'str', 'Code employé créateur'),
        ('updated_at', 'datetime?', 'Date de modification'),
        ('updated_by', 'str?', 'Code employé modificateur'),
    ]
    for attr in supply_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # Table pour SupplyItem
    doc.add_heading('SupplyItem (Ligne d\'approvisionnement)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    supplyitem_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('product_id', 'str', 'FK → Product'),
        ('product_name', 'str?', 'Nom produit (dénormalisé)'),
        ('quantity', 'int', 'Quantité'),
        ('unit_price', 'float', 'Prix unitaire d\'achat'),
        ('total_price', 'float', 'Prix total'),
        ('date_peremption', 'datetime?', 'Date de péremption du lot'),
    ]
    for attr in supplyitem_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # =====================================
    # SECTION 5: STOCK & HISTORIQUE
    # =====================================
    doc.add_heading('5. Stock & Historique', level=1)
    
    # Table pour StockMovement
    doc.add_heading('StockMovement (Mouvement de stock)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    stock_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('product_id', 'str', 'FK → Product'),
        ('product_name', 'str?', 'Nom produit (dénormalisé)'),
        ('movement_type', 'enum', 'initial|supply|sale|return|adjustment'),
        ('movement_quantity', 'int', 'Quantité (+/-)'),
        ('stock_before', 'int', 'Stock avant mouvement'),
        ('stock_after', 'int', 'Stock après mouvement'),
        ('reference_type', 'str?', 'Type de référence'),
        ('reference_id', 'str?', 'ID de la référence'),
        ('notes', 'str?', 'Notes'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date du mouvement'),
        ('created_by', 'str', 'Code employé'),
    ]
    for attr in stock_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_page_break()
    
    # =====================================
    # SECTION 6: AUTRES ENTITÉS
    # =====================================
    doc.add_heading('6. Autres Entités', level=1)
    
    # Table pour Prescription
    doc.add_heading('Prescription (Ordonnance)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    prescription_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('customer_id', 'str', 'FK → Customer'),
        ('doctor_name', 'str', 'Nom du médecin'),
        ('medications', 'List[Dict]', 'Liste des médicaments'),
        ('notes', 'str?', 'Notes'),
        ('status', 'str', 'pending | fulfilled'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('created_at', 'datetime', 'Date de création'),
    ]
    for attr in prescription_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_paragraph()
    
    # Table pour Settings
    doc.add_heading('Settings (Paramètres)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attribut'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    settings_attrs = [
        ('id', 'str (UUID)', 'Identifiant unique'),
        ('tenant_id', 'str', 'ID de l\'agence'),
        ('stock_valuation_method', 'str', 'fifo | lifo | weighted_average'),
        ('currency', 'str', 'Devise (GNF, EUR, USD)'),
        ('pharmacy_name', 'str?', 'Nom de la pharmacie'),
        ('low_stock_threshold', 'int', 'Seuil de stock bas'),
        ('return_delay_days', 'int', 'Délai max pour retours (jours)'),
        ('expiration_alert_days', 'int', 'Alerte péremption (jours)'),
        ('created_at', 'datetime', 'Date de création'),
        ('updated_at', 'datetime', 'Date de modification'),
    ]
    for attr in settings_attrs:
        row = table.add_row().cells
        row[0].text = attr[0]
        row[1].text = attr[1]
        row[2].text = attr[2]
    
    doc.add_page_break()
    
    # =====================================
    # SECTION 7: DIAGRAMME DES RELATIONS
    # =====================================
    doc.add_heading('7. Relations entre les Entités', level=1)
    
    relations_text = """
┌─────────────────────────────────────────────────────────────────────────────┐
│                          DIAGRAMME DE RELATIONS                              │
├─────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│    ┌──────────┐                                                              │
│    │  User    │                                                              │
│    │ (Admin,  │                                                              │
│    │ Pharmac.,│                                                              │
│    │ Caissier)│                                                              │
│    └────┬─────┘                                                              │
│         │ créé par (employee_code)                                           │
│         ▼                                                                    │
│    ┌──────────┐    contient    ┌──────────┐    appartient à   ┌──────────┐  │
│    │ Category │◄───────────────│ Product  │─────────────────►│  Unit    │  │
│    │          │                │          │                   │          │  │
│    │ markup   │                │ stock    │                   │ Boîte,   │  │
│    │ coeff.   │                │ prix     │                   │ Flacon..│  │
│    └──────────┘                │ péremp.  │                   └──────────┘  │
│                                └────┬─────┘                                  │
│                                     │                                        │
│         ┌───────────────────────────┼───────────────────────────┐            │
│         │                           │                           │            │
│         ▼                           ▼                           ▼            │
│    ┌──────────┐              ┌──────────┐              ┌──────────┐         │
│    │  Sale    │              │ Supply   │              │StockMove │         │
│    │          │              │ment      │              │ment      │         │
│    │ VNT-XXX  │              │          │              │          │         │
│    │ items[]  │              │ items[]  │              │ +/- qty  │         │
│    └────┬─────┘              └────┬─────┘              └──────────┘         │
│         │                         │                                          │
│         │ référence               │ provient de                             │
│         ▼                         ▼                                          │
│    ┌──────────┐              ┌──────────┐                                    │
│    │SaleReturn│              │ Supplier │                                    │
│    │          │              │          │                                    │
│    │ RET-XXX  │              │ is_active│                                    │
│    │ reason   │              │          │                                    │
│    └──────────┘              └──────────┘                                    │
│                                                                              │
│    ┌──────────┐              ┌──────────┐              ┌──────────┐         │
│    │ Customer │◄─────────────│Prescrip- │              │ Settings │         │
│    │          │  pour        │  tion    │              │          │         │
│    │          │              │          │              │ devise   │         │
│    │          │              │ status   │              │ seuils   │         │
│    └──────────┘              └──────────┘              └──────────┘         │
│                                                                              │
└─────────────────────────────────────────────────────────────────────────────┘

LÉGENDE:
────────
→  : Relation (FK)
◄─ : Relation inverse
│  : Héritage / Composition

CARDINALITÉS:
─────────────
• User → Sale/Supply/Return : 1..* (un utilisateur peut créer plusieurs)
• Category → Product : 1..* (une catégorie contient plusieurs produits)
• Supplier → Supply : 1..* (un fournisseur peut avoir plusieurs appros)
• Customer → Sale : 1..* (un client peut avoir plusieurs ventes)
• Customer → Prescription : 1..* (un client peut avoir plusieurs ordonnances)
• Sale → SaleReturn : 1..* (une vente peut avoir plusieurs retours)
• Product → StockMovement : 1..* (un produit a plusieurs mouvements)
• Settings : 1 par tenant (singleton par agence)
"""
    
    # Ajouter le diagramme en texte
    para = doc.add_paragraph()
    run = para.add_run(relations_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    
    doc.add_page_break()
    
    # =====================================
    # SECTION 8: RÉSUMÉ
    # =====================================
    doc.add_heading('8. Résumé des Classes', level=1)
    
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'
    hdr = summary_table.rows[0].cells
    hdr[0].text = 'Classe'
    hdr[1].text = 'Description'
    hdr[2].text = 'Multi-tenant'
    hdr[3].text = 'Traçabilité'
    
    summary_data = [
        ('User', 'Utilisateur du système', 'Oui', 'employee_code'),
        ('Product', 'Produit pharmaceutique', 'Oui', 'created_at, updated_at'),
        ('Category', 'Catégorie de produits', 'Oui', 'markup_coefficient'),
        ('Unit', 'Unité de mesure', 'Oui', '-'),
        ('Supplier', 'Fournisseur', 'Oui', 'is_active, updated_by'),
        ('Customer', 'Client', 'Oui', '-'),
        ('Sale', 'Vente', 'Oui', 'sale_number, employee_code'),
        ('SaleReturn', 'Retour de vente', 'Oui', 'return_number, employee_code'),
        ('Supply', 'Approvisionnement', 'Oui', 'created_by, validated_by'),
        ('SupplyItem', 'Ligne d\'appro', 'via Supply', 'date_peremption'),
        ('StockMovement', 'Mouvement de stock', 'Oui', 'created_by'),
        ('Prescription', 'Ordonnance', 'Oui', 'status'),
        ('Settings', 'Paramètres', 'Oui (1/tenant)', 'updated_at'),
    ]
    
    for data in summary_data:
        row = summary_table.add_row().cells
        row[0].text = data[0]
        row[1].text = data[1]
        row[2].text = data[2]
        row[3].text = data[3]
    
    # Sauvegarder le document
    output_path = '/app/DynSoft_Pharma_Diagramme_Classes.docx'
    doc.save(output_path)
    print(f"Document généré: {output_path}")
    return output_path

if __name__ == "__main__":
    create_class_diagram_doc()
